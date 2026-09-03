import docx

for doc_path in [
    'a/BỘ ĐỀ LISTENING/VSTEP LISTENING 1/Btvn 5 - Nghe.docx',
    'a/BỘ ĐỀ LISTENING/VSTEP LISTENING 2/VSTEP LISTENING TEST 2.docx',
    'a/BỘ ĐỀ LISTENING/VSTEP LISTENING 3/VSTEP LISTENING 3.docx'
]:
    doc = docx.Document(doc_path)
    print(f"=== {doc_path} ===")
    print(f"Tables: {len(doc.tables)}")
    for t_idx, tbl in enumerate(doc.tables):
        print(f" Table {t_idx+1}: {len(tbl.rows)} rows, {len(tbl.columns)} cols")
        for r in tbl.rows[:3]:
            row_txt = [c.text.strip().replace('\n', ' ') for c in r.cells]
            print(f"   {row_txt[:4]}")
